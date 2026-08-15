"""Generate reusable software design document templates."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any

DOC_TYPES = {
    "requirements": {
        "label": "需求分析",
        "filename": "requirements-analysis.md",
        "sections": [
            ("1. 背景与目标", "说明业务背景、问题、建设目标和成功指标。"),
            ("2. 范围与边界", "明确本期范围、非本期范围、上下游系统和约束。"),
            ("3. 用户角色", "| 角色 | 职责 | 权限 | 使用场景 |\n|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | 待补充 |"),
            ("4. 业务流程", "```mermaid\nflowchart TD\n    A[开始] --> B[待补充业务步骤]\n    B --> C[结束]\n```"),
            ("5. 功能需求", "| 编号 | 功能 | 描述 | 优先级 | 验收标准 |\n|---|---|---|---|---|\n| FR-001 | 待补充 | 待补充 | P0 | Given / When / Then |"),
            ("6. 非功能需求", "| 类别 | 指标/要求 | 验证方式 |\n|---|---|---|\n| 性能 | 待补充 | 压测/监控 |\n| 安全 | 待补充 | 安全测试 |\n| 可用性 | 待补充 | 演练/检查 |"),
            ("7. 业务规则与数据", "列出计算规则、状态转换、数据保留、数据质量和合规要求。"),
            ("8. 外部接口与依赖", "| 依赖方 | 接口/事件 | 输入输出 | 失败处理 |\n|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | 待补充 |"),
            ("9. 风险、假设与待确认项", "| 类型 | 内容 | 负责人 | 截止时间 | 状态 |\n|---|---|---|---|---|\n| 待确认 | 待补充 | 待补充 | 待补充 | open |"),
            ("10. 需求验收清单", "- [ ] 所有功能需求都有唯一编号和验收标准\n- [ ] 非功能指标可测量\n- [ ] 角色、流程、数据和外部依赖已覆盖\n- [ ] 待确认项已有负责人"),
        ],
    },
    "high_level_design": {
        "label": "概要设计",
        "filename": "high-level-design.md",
        "sections": [
            ("1. 设计目标与约束", "说明设计目标、范围、技术约束、容量假设和关键取舍。"),
            ("2. 总体架构", "```mermaid\nflowchart LR\n    Client[客户端] --> Gateway[接入层]\n    Gateway --> Service[业务服务]\n    Service --> DB[(数据库)]\n    Service --> External[外部依赖]\n```"),
            ("3. 模块划分", "| 模块 | 职责 | 输入 | 输出 | 依赖 |\n|---|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | 待补充 | 待补充 |"),
            ("4. 关键业务流程", "描述主流程、异常流程、幂等策略、事务边界和异步消息。"),
            ("5. 接口总览", "| 接口 | 方法/事件 | 调用方 | 主要数据 | SLA |\n|---|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | 待补充 | 待补充 |"),
            ("6. 数据架构", "说明核心实体、数据归属、读写路径、缓存、搜索和数据同步策略。"),
            ("7. 部署与运行架构", "说明环境、服务实例、网络分区、配置、日志、监控和发布回滚。"),
            ("8. 安全设计", "说明认证、授权、敏感数据、审计、输入校验和密钥管理。"),
            ("9. 非功能设计", "| 维度 | 目标 | 方案 | 验证 |\n|---|---|---|---|\n| 性能 | 待补充 | 待补充 | 待补充 |\n| 可用性 | 待补充 | 待补充 | 待补充 |\n| 可维护性 | 待补充 | 待补充 | 待补充 |"),
            ("10. 关键决策记录", "| 决策 | 选项 | 结论 | 原因 | 日期 |\n|---|---|---|---|---|\n| ADR-001 | 待补充 | 待补充 | 待补充 | 待补充 |"),
        ],
    },
    "database_design": {
        "label": "数据库设计",
        "filename": "database-design.md",
        "sections": [
            ("1. 设计目标与原则", "说明数据边界、规范化/反规范化原则、命名规范、审计和生命周期策略。"),
            ("2. 数据库与容量概览", "| 项目 | 设计值 | 依据 |\n|---|---|---|\n| 数据库类型 | 待补充 | 待补充 |\n| 初始数据量 | 待补充 | 待补充 |\n| 年增长量 | 待补充 | 待补充 |\n| 峰值并发 | 待补充 | 待补充 |"),
            ("3. 实体关系模型", "```mermaid\nerDiagram\n    ENTITY_A ||--o{ ENTITY_B : contains\n    ENTITY_A {\n        string id PK\n    }\n    ENTITY_B {\n        string id PK\n        string entity_a_id FK\n    }\n```"),
            ("4. 表设计", "| 表名 | 中文名 | 说明 | 主键 | 关联 | 数据保留 |\n|---|---|---|---|---|---|\n| t_example | 示例表 | 待补充 | id | 待补充 | 待补充 |"),
            ("5. 字段字典", "| 表名 | 字段 | 类型 | 必填 | 默认值 | 说明 | 脱敏 |\n|---|---|---|---|---|---|---|\n| t_example | id | varchar(64) | 是 | - | 主键 | 否 |"),
            ("6. 索引、约束与一致性", "列出唯一约束、外键、检查约束、索引选择、事务隔离、并发冲突和幂等策略。"),
            ("7. SQL、迁移与初始化", "说明建表脚本、迁移顺序、回滚方案、初始化数据和兼容旧版本的策略。"),
            ("8. 性能与运维", "说明慢查询、分区/分表、备份恢复、归档、监控告警和容量扩展。"),
            ("9. 安全与合规", "说明账号权限、敏感字段、加密、审计、访问控制和数据清理。"),
            ("10. 数据库验收清单", "- [ ] 表、字段、索引和约束均有来源需求\n- [ ] 主键、唯一性、关联和删除策略明确\n- [ ] 迁移、回滚、备份恢复方案可执行\n- [ ] 容量、性能和安全指标可验证"),
        ],
    },
    "detailed_design": {
        "label": "详细设计",
        "filename": "detailed-design.md",
        "sections": [
            ("1. 模块职责与边界", "说明模块职责、调用方、被调用方、输入输出和不变量。"),
            ("2. 类与组件设计", "| 类/组件 | 职责 | 关键属性 | 关键方法 | 依赖 |\n|---|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | 待补充 | 待补充 |"),
            ("3. 时序与交互", "```mermaid\nsequenceDiagram\n    participant C as 调用方\n    participant S as 服务\n    participant R as 仓储\n    C->>S: 请求\n    S->>R: 读写数据\n    R-->>S: 结果\n    S-->>C: 响应\n```"),
            ("4. 接口详细定义", "| 项目 | 内容 |\n|---|---|\n| 请求 | 待补充 |\n| 参数 | 待补充 |\n| 成功响应 | 待补充 |\n| 错误响应 | 待补充 |\n| 权限与幂等 | 待补充 |"),
            ("5. 核心算法与伪代码", "```text\n输入：待补充\n1. 校验输入\n2. 执行业务规则\n3. 持久化或发布事件\n4. 返回结果\n```"),
            ("6. 状态与异常处理", "| 场景/状态 | 处理方式 | 用户提示 | 日志级别 | 重试 |\n|---|---|---|---|---|\n| 待补充 | 待补充 | 待补充 | INFO | 否 |"),
            ("7. 配置、日志与可观测性", "列出配置项、日志字段、指标、链路追踪、告警阈值和敏感信息处理。"),
            ("8. 并发、事务与性能", "说明锁、事务边界、超时、缓存、批处理、限流和资源释放。"),
            ("9. 测试设计", "| 测试层级 | 场景 | 关键断言 | 测试数据 |\n|---|---|---|---|\n| 单元 | 待补充 | 待补充 | 待补充 |\n| 集成 | 待补充 | 待补充 | 待补充 |\n| 接口 | 待补充 | 待补充 | 待补充 |"),
            ("10. 实现验收清单", "- [ ] 需求编号可追溯到模块、接口或数据变更\n- [ ] 正常、异常、边界和并发场景已覆盖\n- [ ] 配置、日志、监控和回滚已定义\n- [ ] 测试用例可执行"),
        ],
    },
}

DEFAULT_TITLES = {key: value["label"] for key, value in DOC_TYPES.items()}
VALID_CONFLICTS = frozenset({"skip", "rename", "overwrite"})
VALID_OUTPUT_FORMATS = frozenset({"auto", "markdown", "md", "docx"})
FIGURE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".svg"})
DOCX_FIGURE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg"})
FIGURE_RE = re.compile(r"!\[(?P<alt>[^]]*)\]\((?P<path>[^)]+)\)")
DOCX_WIDTH_DXA = 9360
DOCX_INDENT_DXA = 120
DOCX_CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def _agent_root() -> Path:
    current = Path(__file__).resolve().parent
    for directory in (current, *current.parents):
        evolve_marker = directory / "evolve"
        if (evolve_marker / "_index.core.toml").is_file() or (evolve_marker / "_index.toml").is_file():
            return directory
    raise RuntimeError("could not locate agent root")


def _load_paths():
    core = _agent_root() / "agent-core"
    if str(core) not in sys.path:
        sys.path.insert(0, str(core))
    from paths import AgentPaths, PathDeniedForWriteError, PathOutOfBoundsError

    return AgentPaths, PathDeniedForWriteError, PathOutOfBoundsError


def _renamed_target(target: Path) -> Path:
    index = 1
    while True:
        candidate = target.with_name(f"{target.stem}-{index}{target.suffix}")
        if not candidate.exists():
            return candidate
        index += 1


def _validate_figures(payload: dict[str, Any], paths, output_format: str) -> list[dict[str, Any]]:
    raw_figures = payload.get("figures", [])
    if raw_figures is None:
        return []
    if not isinstance(raw_figures, list):
        raise ValueError("figures must be an array")
    figures: list[dict[str, Any]] = []
    for index, raw_figure in enumerate(raw_figures):
        if not isinstance(raw_figure, dict):
            raise ValueError(f"figures[{index}] must be an object")
        figure_path = raw_figure.get("path")
        if not isinstance(figure_path, str) or not figure_path.strip():
            raise ValueError(f"figures[{index}].path is required")
        try:
            absolute_path = paths.resolve_under_agent(figure_path, must_exist=True)
        except Exception as exc:
            raise ValueError(f"figures[{index}].path is invalid: {exc}") from exc
        suffix = absolute_path.suffix.lower()
        if suffix not in FIGURE_SUFFIXES:
            raise ValueError(f"figures[{index}].path must use .png, .jpg, .jpeg or .svg")
        if output_format == "docx" and suffix not in DOCX_FIGURE_SUFFIXES:
            raise ValueError("DOCX figures must be PNG or JPEG; render SVG as PNG first")
        caption = raw_figure.get("caption", "")
        alt_text = raw_figure.get("alt_text") or caption or absolute_path.stem
        section = raw_figure.get("section", "")
        if not isinstance(caption, str) or not isinstance(alt_text, str) or not isinstance(section, str):
            raise ValueError(f"figures[{index}].caption, alt_text and section must be strings")
        width_inches = raw_figure.get("width_inches", 6.2)
        try:
            width_inches = float(width_inches)
        except (TypeError, ValueError) as exc:
            raise ValueError(f"figures[{index}].width_inches must be a number") from exc
        if not 1.0 <= width_inches <= 6.5:
            raise ValueError(f"figures[{index}].width_inches must be between 1 and 6.5")
        figures.append({
            "path": paths.to_agent_relative(absolute_path),
            "absolute_path": absolute_path,
            "caption": caption.strip(),
            "alt_text": alt_text.strip(),
            "section": section.strip(),
            "width_inches": width_inches,
        })
    return figures


def _figure_markdown(figure: dict[str, Any]) -> str:
    text = f"![{figure['alt_text']}]({figure['path']})"
    if figure["caption"]:
        text += f"\n\n*{figure['caption']}*"
    return text


def _render(payload: dict[str, Any], doc_type: str, figures: list[dict[str, Any]] | None = None) -> str:
    spec = DOC_TYPES[doc_type]
    project_name = str(payload.get("project_name") or "未命名项目").strip()
    title = str(payload.get("title") or DEFAULT_TITLES[doc_type]).strip()
    summary = str(payload.get("summary") or "待补充").strip()
    figures = figures or []
    rendered_sections = []
    for heading, body in spec["sections"]:
        matching_figures = [figure for figure in figures if figure["section"] == heading]
        figure_text = "\n\n".join(_figure_markdown(figure) for figure in matching_figures)
        rendered_sections.append(f"## {heading}\n\n{body}" + (f"\n\n{figure_text}" if figure_text else ""))
    sections = "\n\n".join(rendered_sections)
    return (
        f"# {title}\n\n"
        f"> 项目：{project_name}\n> 文档类型：{spec['label']}\n> 状态：草稿\n\n"
        f"## 文档说明\n\n{summary}\n\n"
        f"{sections}\n"
    )


def _resolve_output_format(payload: dict[str, Any], target: Path) -> str:
    requested = payload.get("output_format", "auto")
    if not isinstance(requested, str) or requested.strip().lower() not in VALID_OUTPUT_FORMATS:
        raise ValueError(f"output_format must be one of {sorted(VALID_OUTPUT_FORMATS)}")
    requested = requested.strip().lower()
    if requested == "auto":
        return "docx" if target.suffix.lower() == ".docx" else "markdown"
    if requested == "docx" and target.suffix.lower() != ".docx":
        raise ValueError("output_format=docx requires a .docx path")
    if requested in {"markdown", "md"} and target.suffix.lower() == ".docx":
        raise ValueError("Markdown output cannot use a .docx path")
    return "docx" if requested == "docx" else "markdown"


def _set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_docx_styles(document) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2933")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    metadata = document.styles.add_style("Design Metadata", WD_STYLE_TYPE.PARAGRAPH)
    metadata.font.name = "Calibri"
    metadata.font.size = Pt(9.5)
    metadata.font.color.rgb = RGBColor.from_string("667085")
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.0

    code = document.styles.add_style("Design Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string("344054")
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    caption = document.styles.add_style("Design Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string("667085")
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = 1


def _set_cell_margins(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in DOCX_CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_header_row(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    row_properties = row._tr.get_or_add_trPr()
    header = row_properties.find(qn("w:tblHeader"))
    if header is None:
        row_properties.append(OxmlElement("w:tblHeader"))


def _set_table_geometry(table, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(DOCX_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), str(DOCX_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")
    table_layout = table_properties.first_child_found_in("w:tblLayout")
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_properties.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = tc_pr.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[column_index]))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _table_widths(column_count: int) -> list[int]:
    patterns = {
        2: [2700, 6660],
        3: [1800, 2800, 4760],
        4: [1200, 2200, 2800, 3160],
        5: [1100, 1900, 2500, 1900, 1960],
        6: [900, 1500, 2200, 1500, 1500, 1760],
        7: [850, 1350, 1900, 1450, 1250, 1250, 1310],
    }
    if column_count in patterns:
        return patterns[column_count]
    base_width, remainder = divmod(DOCX_WIDTH_DXA, column_count)
    return [base_width + (1 if index < remainder else 0) for index in range(column_count)]


def _split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _add_docx_table(document, rows: list[list[str]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_header_row(table.rows[0])
    for row_index, row_data in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = row_data[column_index] if column_index < len(row_data) else ""
            _set_cell_shading(cell, "E8EEF5" if row_index == 0 else "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_run_font(run, size=9.5, color="1F2933", bold=row_index == 0)
    _set_table_geometry(table, _table_widths(column_count))
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_docx_figure(document, figure: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(figure["absolute_path"]), width=Inches(figure["width_inches"]))
    shape._inline.docPr.set("descr", figure["alt_text"])
    shape._inline.docPr.set("title", figure["caption"] or figure["alt_text"])
    if figure["caption"]:
        document.add_paragraph(figure["caption"], style="Design Figure Caption")


def _render_docx(markdown: str, target: Path, payload: dict[str, Any], doc_type: str, figures: list[dict[str, Any]] | None = None) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("DOCX 输出需要安装 python-docx，请运行 python -m pip install -r requirements.txt") from exc

    document = Document()
    _configure_docx_styles(document)
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"{payload.get('project_name', '未命名项目')} · {DOC_TYPES[doc_type]['label']}"
    header.style = document.styles["Design Metadata"]
    footer = section.footer.paragraphs[0]
    footer.text = "design_document · 草稿"
    footer.style = document.styles["Design Metadata"]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    figure_map = {figure["path"]: figure for figure in (figures or [])}
    lines = markdown.splitlines()
    line_index = 0
    while line_index < len(lines):
        line = lines[line_index]
        if not line.strip():
            line_index += 1
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            run = paragraph.add_run(line[2:].strip())
            _set_run_font(run, size=24, color="0B2545", bold=True)
            line_index += 1
            continue
        if line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            line_index += 1
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            line_index += 1
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="Design Metadata")
            run = paragraph.add_run(line[2:].strip())
            _set_run_font(run, size=9.5, color="667085")
            line_index += 1
            continue
        image_match = FIGURE_RE.fullmatch(line.strip())
        if image_match:
            figure = figure_map.get(image_match.group("path"))
            if figure is not None:
                _add_docx_figure(document, figure)
            line_index += 1
            if line_index < len(lines) and lines[line_index].strip().startswith("*") and lines[line_index].strip().endswith("*"):
                line_index += 1
            continue
        if line.startswith("```"):
            line_index += 1
            while line_index < len(lines) and not lines[line_index].startswith("```"):
                paragraph = document.add_paragraph(style="Design Code")
                run = paragraph.add_run(lines[line_index])
                _set_run_font(run, name="Consolas", size=9, color="344054")
                line_index += 1
            line_index += 1
            continue
        if line.startswith("|") and line_index + 1 < len(lines) and _is_table_separator(lines[line_index + 1]):
            rows = [_split_table_row(line)]
            line_index += 2
            while line_index < len(lines) and lines[line_index].startswith("|"):
                rows.append(_split_table_row(lines[line_index]))
                line_index += 1
            _add_docx_table(document, rows)
            continue
        if line.startswith("- "):
            text = line[2:].strip()
            if text.startswith("[ ] "):
                text = "☐ " + text[4:]
            elif text.startswith("[x] "):
                text = "☒ " + text[4:]
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(text)
            _set_run_font(run, size=11, color="1F2933")
            line_index += 1
            continue
        paragraph = document.add_paragraph(line.strip())
        for run in paragraph.runs:
            _set_run_font(run, size=11, color="1F2933")
        line_index += 1
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def run_design_document(payload: dict[str, Any]) -> dict[str, Any]:
    AgentPaths, PathDeniedForWriteError, PathOutOfBoundsError = _load_paths()
    paths = AgentPaths.discover(start=_agent_root())
    path_arg = payload.get("path")
    doc_type = payload.get("doc_type")
    if not isinstance(path_arg, str) or not path_arg.strip():
        return {"ok": False, "error": "path is required"}
    if not isinstance(doc_type, str) or doc_type not in DOC_TYPES:
        return {"ok": False, "error": f"doc_type must be one of {sorted(DOC_TYPES)}"}
    on_conflict = payload.get("on_conflict", "skip")
    if not isinstance(on_conflict, str) or on_conflict.strip().lower() not in VALID_CONFLICTS:
        return {"ok": False, "error": f"on_conflict must be one of {sorted(VALID_CONFLICTS)}"}
    on_conflict = on_conflict.strip().lower()
    try:
        target = paths.resolve_under_agent_for_write(path_arg, must_exist=False)
    except (PathOutOfBoundsError, PathDeniedForWriteError, TypeError, ValueError) as exc:
        return {"ok": False, "error": str(exc)}
    try:
        output_format = _resolve_output_format(payload, target)
    except ValueError as exc:
        return {"ok": False, "error": str(exc)}
    try:
        figures = _validate_figures(payload, paths, output_format)
    except ValueError as exc:
        return {"ok": False, "error": str(exc)}
    section_count = len(DOC_TYPES[doc_type]["sections"])
    if target.exists() and on_conflict == "skip":
        return {"ok": True, "dry_run": bool(payload.get("dry_run", False)), "path": paths.to_agent_relative(target), "written": False, "doc_type": doc_type, "format": output_format, "sections": section_count, "figures": len(figures)}
    if target.exists() and on_conflict == "rename":
        target = _renamed_target(target)
    relative_path = paths.to_agent_relative(target)
    if bool(payload.get("dry_run", False)):
        return {"ok": True, "dry_run": True, "path": relative_path, "written": False, "doc_type": doc_type, "format": output_format, "sections": section_count, "figures": len(figures)}
    try:
        markdown = _render(payload, doc_type, figures)
        if output_format == "docx":
            _render_docx(markdown, target, payload, doc_type, figures)
        else:
            target.parent.mkdir(parents=True, exist_ok=True)
            from evolve_tool_io import write_utf8_text

            write_utf8_text(target, markdown)
    except (OSError, RuntimeError) as exc:
        return {"ok": False, "error": str(exc)}
    return {"ok": True, "path": relative_path, "written": True, "doc_type": doc_type, "format": output_format, "sections": section_count, "figures": len(figures)}


def main() -> None:
    core = _agent_root() / "agent-core"
    if str(core) not in sys.path:
        sys.path.insert(0, str(core))
    from evolve_tool_io import run_tool_main

    run_tool_main(run_design_document)


def _demo() -> None:
    for doc_type, spec in DOC_TYPES.items():
        rendered = _render({"project_name": "demo", "summary": "summary"}, doc_type)
        assert rendered.startswith("# ") and spec["label"] in rendered
        assert len(spec["sections"]) >= 8
    print(json.dumps({"ok": True, "doc_types": sorted(DOC_TYPES), "templates": len(DOC_TYPES), "formats": ["markdown", "docx"]}, ensure_ascii=False))


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "demo":
        _demo()
    else:
        main()
